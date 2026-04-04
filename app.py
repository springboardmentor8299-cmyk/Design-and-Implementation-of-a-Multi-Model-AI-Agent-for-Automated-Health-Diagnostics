from click import style
import streamlit as st
import json
import os
import pandas as pd
import requests  
from processing.synthesis_engine import synthesize_findings
from processing.recommendation_engine import generate_recommendations
from processing.risk_analyzer import analyze_risk
from processing.comparator import compare_with_ranges
from scanner.ocr_engine import scan_pdf, scan_image
from processing.extractor import extract_parameters_from_text
from api.medical_api import fetch_reference_ranges
from evaluation.evaluate import evaluate_single_report
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from PIL import Image, ImageDraw

st.set_page_config(page_title="Blood Report Analyzer", layout="wide")

st.markdown("""
<style>

/* 🔥 BLACK BACKGROUND */
[data-testid="stAppViewContainer"] {
    background-color: #0e1117;
}

/* Title */
.title {
    text-align: center;
    color: white;
}

/* Cards (dark glass style) */
.card {
    padding: 20px;
    border-radius: 15px;
    text-align: center;
    margin: 10px;
    background: rgba(255,255,255,0.05);
    color: white;
    border: 1px solid rgba(255,255,255,0.1);
    backdrop-filter: blur(8px);
}

/* Status colors */
.normal { background: rgba(76,175,80,0.2); }
.low { background: rgba(255,193,7,0.2); }
.high { background: rgba(244,67,54,0.2); }

/* Buttons */
div.stButton > button {
    width: 100%;
    border-radius: 25px;
    height: 45px;
    background: linear-gradient(90deg,#00b4db,#0083b0);
    color: white;
}

/* Fix default text color */
html, body, [class*="css"] {
    color: white;
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<h1 class="title">🩺 Blood Report Analyzer</h1>
<p class="title">AI-Powered Medical Dashboard</p>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Upload Report", type=["pdf","png","jpg","jpeg","json"])

if "tab" not in st.session_state:
    st.session_state.tab = "parameters"

c1,c2,c3,c4,c5 = st.columns([2,1,1,1,2])

with c2:
    if st.button("📊 Parameters"):
        st.session_state.tab="parameters"
with c3:
    if st.button("🧠 Risks"):
        st.session_state.tab="risks"
with c4:
    if st.button("💡 AI Insights"):
        st.session_state.tab="insights"


if uploaded_file:

    ext = uploaded_file.name.split(".")[-1].lower()
    temp_file = f"temp.{ext}"

    with open(temp_file,"wb") as f:
        f.write(uploaded_file.read())

    if ext=="pdf":
        scanned_text=scan_pdf(temp_file)
    elif ext in ["png","jpg","jpeg"]:
        scanned_text=scan_image(temp_file)
    else:
        scanned_text=json.load(open(temp_file))

    if isinstance(scanned_text,dict):
        df=pd.DataFrame(list(scanned_text.items()),columns=["Parameter","Value"])
    else:
        df=extract_parameters_from_text(scanned_text)

    if df.empty:
        st.error("No parameters detected")
        st.stop()

    response = requests.post(
        "http://127.0.0.1:8000/analyze",
        json={"parameters": df.to_dict(orient="records")}
    )

    df = pd.DataFrame(response.json())
    risk=analyze_risk(df)
    # ✅ Milestone 4: Create Final Report Object
    final_report = {
        "parameters": df.to_dict(orient="records"),
        "risk_analysis": risk,
        "summary": None,
        "recommendations": None,
        "disclaimer": "⚠ This is AI-generated and not a substitute for medical advice."
    }

    if st.session_state.tab=="parameters":

        cols=st.columns(4)
        for i,row in df.iterrows():
            status=row["Status"]
            style="normal" if status=="Normal" else "low" if status=="Low" else "high"
            icon="🟢" if status=="Normal" else "🟡" if status=="Low" else "🔴"

            cols[i%4].markdown(f"""
            <div class="card {style}">
                <h4>{icon} {row['Parameter']}</h4>
                <h2>{row['Value']}</h2>
                <p>{status}</p>
            </div>
            """,unsafe_allow_html=True)

        st.dataframe(df)


    elif st.session_state.tab=="risks":

        st.subheader("🧠 Risk Analysis")

        high=[k for k,v in risk.items() if v["level"]=="HIGH RISK"]
        if high:
            st.error("🚨 Immediate Attention Required: "+", ".join(high))

        if "Diabetes Risk" in high and "Cardiovascular Risk" in high:
            st.error("⚠ Combined metabolic risk detected")

        sorted_risks=sorted(risk.items(), key=lambda x: x[1]["score"], reverse=True)

        st.subheader("🔝 Risk Priority Ranking")
        for i,r in enumerate(sorted_risks):
            st.write(f"{i+1}. {r[0]} → {r[1]['level']}")

        cols=st.columns(4)
        for i,(k,v) in enumerate(risk.items()):

            level=v["level"]
            score=v["score"]
            style="normal" if "LOW" in level else "low" if "MODERATE" in level else "high"
            icon="🟢" if "LOW" in level else "🟡" if "MODERATE" in level else "🔴"

            reason=", ".join(v["reasons"]) if v["reasons"] else "No major issues"

            cols[i%4].markdown(f"""
            <div class="card {style}">
                <h4>{icon} {k}</h4>
                <h3>{level}</h3>
                <p>Score: {score}</p>
                <p>{reason}</p>
            </div>
            """,unsafe_allow_html=True)

            st.progress(min(score*15,100))

            if v["reasons"]:
                st.caption("📌 Triggered by: " + ", ".join(v["reasons"]))

            st.write(f"{k} is {level.lower()} due to {reason.lower()}.")

    elif st.session_state.tab=="insights":

        st.subheader("💡 AI Health Insights")

        param_results={row["Parameter"].lower():row["Status"].lower() for _,row in df.iterrows()}
        normalized_risk={k:v["level"].lower().replace(" risk","") for k,v in risk.items()}

        summary=synthesize_findings(param_results,normalized_risk)
        recs=generate_recommendations(param_results,normalized_risk)
        # ✅ Store in final report
        final_report["summary"] = summary
        final_report["recommendations"] = recs

        abnormal=len([v for v in param_results.values() if v!="normal"])
        total=len(param_results)
        score=int((1-abnormal/total)*100)

        st.metric("💖 Health Score",f"{score}%")
        st.progress(score)

        if summary and summary.strip():
            st.subheader("🧾 AI Summary")

            lines = [line.strip() for line in summary.split(".") if line.strip()]

            for line in lines:
                l = line.lower()

                if "high" in l:
                    style = "high"
                elif "low" in l:
                    style = "low"
                else:
                    style = "normal"

                st.markdown(f"<div class='card {style}'>• {line}</div>", unsafe_allow_html=True)


        st.subheader("📊 Health Status")
        st.write("Excellent" if score>90 else "Moderate" if score>70 else "Needs Attention")

        st.subheader("⏳ Future Risk")
        st.write("Stable" if abnormal==0 else "Monitor closely")
        st.subheader("🤖 Ask Your Report")
        q = st.text_input("Ask something")

        if q:
            try:
                response = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": "Bearer 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",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "http://localhost",
                        "X-Title": "Blood Report Analyzer"
                    },
                    json={
                        "model": "openrouter/auto",
                        "messages": [
                            {
                                "role": "system",
                                "content": """You are a medical assistant analyzing blood reports.

                                                STRICT INSTRUCTION:
                                                Answer ONLY about the parameter or topic asked in the question.

                                                - If the user asks about WBC → talk ONLY about WBC
                                                - If the user asks about food → suggest food related to that condition
                                                - Do NOT include unrelated parameters
                                                - Do NOT summarize the full report unless asked

                                                You can give detailed answers, but keep them focused on the question.

                                                Be practical and helpful."""
                            },
                            {
                                "role": "user",
                                "content": f"Report: {param_results}\nQuestion: {q}"
                            }
                        ]
                    }                 
                )

                data = response.json()

                if "choices" in data:
                    answer = data["choices"][0]["message"]["content"]
                    st.write(answer)
                else:
                    st.error(f"API Error: {data}")

            except Exception as e:
                st.error(f"Error: {e}")

       


        report_text = f"""
        Health Score: {score}

        Summary:
        {final_report["summary"]}

        Risks:
        {final_report["risk_analysis"]}

        Recommendations:
        {final_report["recommendations"]}

        Disclaimer:
        {final_report["disclaimer"]}
        """


        st.subheader("🚨 Key Issues")
        for k,v in param_results.items():
            if v!="normal":
                st.markdown(f"<div class='card high'>⚠ {k.upper()} is {v}</div>", unsafe_allow_html=True)

        if all(v=="normal" for v in param_results.values()):
            st.success("🎉 Excellent Health!")

        st.subheader("🧠 AI Explanation")
        for k,v in param_results.items():
            if v=="high":
                st.write(f"{k} is above normal range.")
            elif v=="low":
                st.write(f"{k} is below normal range.")

        st.subheader("💡 Recommendations")
        for cat,items in recs.items():
            for item in items:
                st.markdown(f"<div class='card normal'>✔ {item}</div>", unsafe_allow_html=True)

        st.subheader("🧭 Action Plan")

        actions = []

       
        for param, status in param_results.items():
            if status == "high":
                actions.append(f"🔴 Reduce {param.upper()} levels – adjust diet & monitor regularly")
            elif status == "low":
                actions.append(f"🟡 Improve {param.upper()} – include nutrient-rich foods")

        actions.append("💧 Drink 2–3 liters of water daily")
        actions.append("🥗 Eat balanced meals (fruits, vegetables, protein)")
        actions.append("🏃 Exercise at least 30 minutes daily")
        actions.append("😴 Sleep 7–8 hours regularly")
        actions.append("🧘 Manage stress (meditation / relaxation)")

        for act in actions:
            st.markdown(f"<div class='card normal'>✔ {act}</div>", unsafe_allow_html=True)


        st.subheader("⏳ Suggested Routine")
        st.markdown("""
        <div class='card normal'>
        🌅 Morning: Warm water + light exercise  
        🍳 Breakfast: Protein-rich meal  
        🏃 Afternoon: Stay active + hydrate  
        🍽️ Evening: Light dinner + walk  
        🌙 Night: Sleep early (7–8 hrs)
        </div>
        """, unsafe_allow_html=True)
     
        st.subheader("👤 Patient Details")

        name = st.text_input("Name")
        age = st.number_input("Age", min_value=1, max_value=120, step=1)
        gender = st.selectbox("Gender", ["Male", "Female", "Other"])
        st.subheader("📄 Final Health Report")
        st.markdown(f"""
        <div style="
            padding:25px;
            border-radius:15px;
            background:#111827;
            color:white;
            font-family:Arial;
            border:1px solid rgba(255,255,255,0.1);
        ">

        <h2 style="text-align:center;">🩺 Health Report</h2>

        <hr>

        <b>👤 Name:</b> {name if name else "Not Provided"} <br>
        <b>🎂 Age:</b> {age} <br>
        <b>⚧ Gender:</b> {gender} <br>

        <hr>

        <h3>🧾 Summary</h3>
        <p>{final_report["summary"]}</p>

        <h3>⚠ Risk Analysis</h3>
        <ul>
        {"".join([f"<li>{k}: {v['level']}</li>" for k,v in final_report["risk_analysis"].items()])}
        </ul>

        <h3>💡 Recommendations</h3>
        <ul>
        {"".join([f"<li>{item}</li>" for cat in final_report["recommendations"] for item in final_report["recommendations"][cat]])}
        </ul>

        <hr>

        <p style="color:#ff6b6b;"><b>{final_report["disclaimer"]}</b></p>

        </div>
        """, unsafe_allow_html=True)

        st.subheader("⬇ Download Report")

        file_type = st.selectbox("Select Format", ["PDF", "DOCX", "IMAGE"])

        if st.button("Download Report"):

            risk_text = "\n".join([
                f"{k}: {v['level']} ({', '.join(v['reasons'])})" if v["reasons"] else f"{k}: {v['level']}"
                for k, v in final_report["risk_analysis"].items()
            ])

            rec_text = "\n".join([
                f"- {item}"
                for cat in final_report["recommendations"]
                for item in final_report["recommendations"][cat]
            ]) if any(final_report["recommendations"].values()) else "No recommendations available"

            content = f"""
            Name: {name}
            Age: {age}
            Gender: {gender}

            Summary:
            {final_report["summary"]}

            Risks:
            {risk_text}

            Recommendations:
            {rec_text}

            Disclaimer:
            {final_report["disclaimer"]}
            """

            if file_type == "PDF":
                pdf_path = "report.pdf"
                doc = SimpleDocTemplate(pdf_path)
                styles = getSampleStyleSheet()

                elements = []
                for line in content.split("\n"):
                    elements.append(Paragraph(line, styles["Normal"]))
                    elements.append(Spacer(1,10))

                doc.build(elements)

                with open(pdf_path, "rb") as f:
                    st.download_button("Download PDF", f, file_name="report.pdf")

            
            elif file_type == "DOCX":
                doc = Document()
                doc.add_heading("Health Report", 0)

                for line in content.split("\n"):
                    doc.add_paragraph(line)

                doc.save("report.docx")

                with open("report.docx", "rb") as f:
                    st.download_button("Download DOCX", f, file_name="report.docx")

         
            elif file_type == "IMAGE":
                img = Image.new("RGB", (800, 1000), color="white")
                d = ImageDraw.Draw(img)

                y = 20
                for line in content.split("\n"):
                    d.text((20, y), line, fill="black")
                    y += 30

                img.save("report.png")

                with open("report.png", "rb") as f:
                    st.download_button("Download Image", f, file_name="report.png")
        st.subheader("🚨 Priority Focus")

        critical = [k for k,v in param_results.items() if v != "normal"]

        if critical:
            st.warning("Focus on: " + ", ".join([c.upper() for c in critical]))
        else:
            st.success("No major issues detected 🎉")

        st.warning("⚠ This is not medical advice")

    name=os.path.splitext(uploaded_file.name)[0]
    gt=f"evaluation/ground_truth/{name}.json"

    if os.path.exists(gt):
        gt_data=json.load(open(gt))
        e,c=evaluate_single_report(df,gt_data)
        st.metric("Accuracy",f"{round((e+c)/2,2)}%")