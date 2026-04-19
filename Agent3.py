import streamlit as st
import pandas as pd
import google.generativeai as genai
from PIL import Image, ImageDraw
import io
import pdfplumber
from docx import Document
from fpdf import FPDF
import datetime

# --- 1. CONFIGURATION (Must be the very first Streamlit command) ---
st.set_page_config(page_title="HemoPulseAI", layout="wide", page_icon="🩸")

# --- 2. API SETUP ---
# Replace with your actual API Key
GEMINI_API_KEY = "YOUR_API_KEY_HERE"
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# --- 3. UI STYLING ---
st.markdown("""
    <style>
    .stApp { background-color: #FFFFFF; }
    .main-header { color: #D32F2F; font-weight: 800; font-size: 2.5rem; }
    .stButton>button {
        background-color: #E8F4F8; color: #000000; border-radius: 8px;
        font-weight: bold; border: 1px solid #B2EBF2; transition: 0.3s;
    }
    .stButton>button:hover { background-color: #B2EBF2; border-color: #00BCD4; }
    .specialist-card { 
        padding: 15px; border-radius: 10px; border-left: 5px solid #00BCD4; 
        background-color: #F0F9FA; margin-bottom: 20px;
    }
    </style>
    """, unsafe_allow_html=True)


# --- 4. CORE LOGIC ---

def extract_text(file_bytes, file_type):
    """Handles text extraction from PDFs and Excel sheets."""
    try:
        if file_type == "application/pdf":
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(io.BytesIO(file_bytes))
            return df.to_string()
    except Exception as e:
        return f"Error reading file: {e}"
    return "Unsupported text format."


def analyze_report(file_bytes, file_type, p_info):
    """The 'Brain' of the app: Generates specialized medical advice."""

    # Detailed prompt focusing on the specialist categories requested
    prompt = f"""
    Act as HemoPulseAI, a senior clinical consultant. Analyze this lab report for a {p_info['age']}-year-old {p_info['gender']}.

    1. **Summary Table**: Provide a [Parameter | Result | Reference Range | Status].
    2. **Primary Diagnostic**: Identify the most likely clinical condition based on abnormalities.

    3. **👨‍⚕️ SPECIALIST ADVICE**: 
       - Which specific specialist should they see (e.g., Cardiologist, Hematologist)?
       - List 3 critical questions the patient should ask during the appointment.

    4. **🥗 TARGETED DIET**: 
       - List 3 foods to strictly avoid and 3 foods to increase based on these specific biomarkers.

    5. **🏃 LIFESTYLE CHANGES**: 
       - Suggest 3 actionable habits (e.g., sleep, exercise type, stress management) tailored to these results.

    6. **📉 VITALS TO BALANCE**: 
       - Identify which vitals (e.g., BP, Heart Rate, SpO2) need monitoring and what the 'Target Range' should be.

    Format the response using clean Markdown with bold headers and bullet points.
    """

    if "image" in file_type:
        img = Image.open(io.BytesIO(file_bytes))
        response = model.generate_content([prompt, img])
    else:
        text_content = extract_text(file_bytes, file_type)
        response = model.generate_content([prompt, f"Data:\n{text_content}"])

    return response.text


# --- 5. EXPORT FUNCTIONS ---

def create_pdf(text, p_info):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="HemoPulseAI Clinical Analysis", ln=True, align='C')
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt=f"Patient: {p_info['name']} | Age: {p_info['age']} | Gender: {p_info['gender']}", ln=True,
             align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 6, text.encode('latin-1', 'replace').decode('latin-1'))
    return pdf.output(dest='S').encode('latin-1')


def create_docx(text, p_info):
    doc = Document()
    doc.add_heading('HemoPulseAI Specialist Report', 0)
    doc.add_paragraph(f"Name: {p_info['name']} | Age: {p_info['age']} | Gender: {p_info['gender']}")
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 6. MAIN UI ---

st.markdown('<h1 class="main-header">🩸 HemoPulseAI</h1>', unsafe_allow_html=True)
st.write("Advanced Lab Interpretation with Specialist Routing")
st.divider()

col1, col2 = st.columns([1, 2], gap="large")

with col1:
    st.subheader("👤 Patient Profile")
    name = st.text_input("Full Name", placeholder="e.g. John Doe")
    age = st.number_input("Age", 0, 120, 30)
    gender = st.selectbox("Gender", ["Male", "Female", "Non-binary"])

    st.subheader("📂 Upload Report")
    file = st.file_uploader("Drop PDF, Image, or Excel", type=["pdf", "png", "jpg", "jpeg", "xlsx"])

    if file and st.button("🔍 Generate Specialist Advice"):
        with st.spinner("Analyzing biomarkers and cross-referencing specialist databases..."):
            p_info = {"name": name, "age": age, "gender": gender}
            report_analysis = analyze_report(file.getvalue(), file.type, p_info)
            st.session_state['analysis'] = report_analysis

with col2:
    st.subheader("🔍 Clinical Interpretation")
    if 'analysis' in st.session_state:
        # Highlight the specialist aspect
        st.markdown(
            '<div class="specialist-card"><b>💡 Pro Tip:</b> Take the "Questions for your Doctor" section below to your next appointment.</div>',
            unsafe_allow_html=True)

        st.markdown(st.session_state['analysis'])

        st.divider()
        st.subheader("📥 Export Final Report")

        if name:
            fmt = st.radio("Select Format", ["PDF", "Word (DOCX)"], horizontal=True)
            p_info = {"name": name, "age": age, "gender": gender}

            if fmt == "PDF":
                btn_data = create_pdf(st.session_state['analysis'], p_info)
                st.download_button("Download PDF", btn_data, f"{name}_Analysis.pdf")
            else:
                btn_data = create_docx(st.session_state['analysis'], p_info)
                st.download_button("Download DOCX", btn_data, f"{name}_Analysis.docx")
        else:
            st.warning("Please enter patient name to enable downloads.")
    else:
        st.info("Upload a report on the left to see the AI specialist's findings.")

st.divider()
st.caption(
    "⚠️ **Disclaimer:** HemoPulseAI is an AI-assisted tool for educational purposes. It does not provide medical diagnoses. Always verify findings with a licensed healthcare provider.")